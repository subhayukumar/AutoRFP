import csv
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup
import re
import os
import markdown
import base64
from docx.enum.text import WD_BREAK
import requests
from urllib import parse
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Image
from io import BytesIO
from datetime import date


def download_image(url: str):
    """Download image from url, return file name"""
    data = requests.get(url).content
    file_name = parse.urlparse(url).path.split("/")[-1]
    with open(file_name, "wb") as f:
        f.write(data)
    return file_name



FFT_LOGO_PATH = download_image("https://i.postimg.cc/v8r5hdzg/fft-logo.jpg")


def image_to_base64(image_path):
    with open(image_path, "rb") as image_file:
        return base64.b64encode(image_file.read()).decode()



def add_line_breaks_between_list_items(md_text):
    md_text = re.sub(r'(\n\d+\. .+)(\n\d+\. .+)', r'\1\n\2', md_text)
    md_text = re.sub(r'(\n  - .+)(\n  - .+)', r'\1\n\2', md_text)
    return md_text


def extract_project_name_from_markdown(md_text):
    # Using regex to find the pattern of the project name in Markdown
    match = re.search(r'^### PROJECT NAME\s*\n(.+)$', md_text, re.MULTILINE)
    if match:
        return match.group(1).strip()
    return None


def add_title_page_to_word(doc, project_name, logo_path):
    # print("document: ", doc.text)
   
    # logo centrally aligned
    logo_paragraph = doc.add_paragraph()
    logo_run = logo_paragraph.add_run()
    logo_run.add_picture(logo_path, width=Inches(1.5))
    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
   
    # title centrally aligned
    title = doc.add_paragraph()
    title_run = title.add_run(project_name)
    title_run.bold = True
    title_run.font.size = Pt(36)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER


    # versioning table
    table = doc.add_table(rows=2, cols=2)  
    table.autofit = True


    # Header cells
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Version'
    hdr_cells[1].text = 'Date'


    # First row cells
    today = date.today()
    row1_cells = table.rows[1].cells
    row1_cells[0].text = '1.0'
    row1_cells[1].text = "dd/mm/yyyy"
    # row1_cells[1].text = today.strftime("%B %d, %Y")


    # New page
    section_break = doc.add_paragraph()
    section_break.add_run()
    section_break.add_run().add_break(WD_BREAK.PAGE)  # Adding page break



def add_footer_to_word(doc, company_name):
    # company name to the footer
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.text = company_name


def add_title_page_to_pdf(html_text, project_name, logo_path):
    print(os.system("pwd"))
    print(os.system("ls"))
    print(os.system("ls utilities"))
    logo_base64 = image_to_base64(logo_path)
    title_page_html = f"""
    <div style='text-align:center'>
        <h1 style='font-size:36px; font-weight:bold;'>{project_name}</h1>
        <img src='data:image/jpeg;base64,{logo_base64}' width='150'>
        <table style='width:50%; margin-top:20px; margin-left:auto; margin-right:auto; border-collapse: collapse; border: 1px solid black;'>
            <thead>
                <tr>
                    <th style='border: 1px solid black; padding: 8px;'>Version</th>
                    <th style='border: 1px solid black; padding: 8px;'>Date</th>
                </tr>
            </thead>
            <tbody>
                <tr>
                    <td style='border: 1px solid black; padding: 8px;'>1.0</td>
                    <td style='border: 1px solid black; padding: 8px;'>dd/mm/yyyy</td>
                </tr>
            </tbody>
        </table>
    </div>
    <div style='page-break-after: always;'></div>  <!-- Page break -->
    """
    return title_page_html + html_text


def convert_markdown_to_pdf_and_word(markdown_file_path, pdf_file_path, word_file_path):
    os.makedirs(os.path.dirname(pdf_file_path), exist_ok=True)
    os.makedirs(os.path.dirname(word_file_path), exist_ok=True)
   


    with open(markdown_file_path, 'r', encoding='utf-8') as f:
        md_text = f.read()
   


    project_name = extract_project_name_from_markdown(md_text)
    html_text = markdown.markdown(md_text)
    # print(html_text)


    doc = Document()


    # docx
    if project_name:
        add_title_page_to_word(doc, project_name, FFT_LOGO_PATH)


    soup = BeautifulSoup(html_text, 'html.parser')  
   


    for element in soup:
        if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            # print("found  h1 !!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!")
            p = doc.add_paragraph()
            run = p.add_run(element.text)
            run.bold = True
            run.font.size = Pt(16 if element.name in ['h1', 'h2'] else 14)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT


        elif element.name == 'p':
            img_tag = element.find('img')            
            if img_tag:
                # Extract image URL
                image_url = img_tag['src']
                # Add a paragraph with centered alignment
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Add the image to the paragraph
                run = p.add_run()
                run.add_picture(image_url, width=Inches(4.5))
            else:
                # Add other paragraphs with justified alignment
                p = doc.add_paragraph(element.text)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


        elif element.name == 'ul':
            for item in element.find_all('li'):
                p = doc.add_paragraph()
                p.add_run('• ')
                p.add_run(item.text)
        elif element.name == 'ol':
            for idx, item in enumerate(element.find_all('li'), 1):
                p = doc.add_paragraph()
                p.add_run(f'{idx}. {item.text}')


    print(doc, "generated")
    doc.save(word_file_path)

# def convert_markdown_to_pdf_and_csv(markdown_file_path, excel_file_path):
#     with open(markdown_file_path, mode='r', encoding='utf-8') as file:
#         data_str = file.read()

#     rows = []
#     backend_total = 0
#     frontend_total = 0

#     for row in data_str.strip().split("\n"):
#         columns = row.split(";")
#         if len(columns) >= 6:  
#             if columns[4].isdigit():
#                 backend_total += int(columns[6])  # Sum backend times
#             if columns[5].isdigit():
#                 frontend_total += int(columns[7])  # Sum frontend times
#         rows.append(columns)

#     rows.append([""] * 6 + [f"Total backend days: {backend_total}", f"Total frontend days: {frontend_total}"])


#     with open(excel_file_path, mode='w', newline='', encoding='utf-8') as file:
#         csv_writer = csv.writer(file)
#         csv_writer.writerows(rows)

#     print(f"CSV file has been created at {excel_file_path} with totals.")

def convert_markdown_to_pdf_and_csv(markdown_file_path, excel_file_path):
    with open(markdown_file_path, 'r') as file:
        data_str = file.read()
        
    lines = data_str.splitlines()

    total_backend_time = 0
    total_frontend_time = 0

    with open(excel_file_path, mode="w", newline="") as file:
        writer = csv.writer(file)

        for line in lines:
            row_data = line.split(';')
            
            # Check if the row has the expected number of columns (6 columns: S.No, Module Name, Subtask Detail, Description, Backend, Frontend)
            if len(row_data) >= 6:
                try:
                    # Extract backend and frontend time
                    backend_time = int(row_data[4])
                    frontend_time = int(row_data[5])
                    
                    # Add to the totals
                    total_backend_time += backend_time
                    total_frontend_time += frontend_time
                except ValueError:
                    # If conversion to int fails, treat as 0
                    backend_time = 0
                    frontend_time = 0
                
                # Write the row data to the CSV file
                writer.writerow(row_data)
            else:
                print(f"Skipping invalid row: {line}")
        
        writer.writerow(["", "", "", "", f"total backend days: {total_backend_time}", f"total frontend days: {total_frontend_time}"])

    print(f"CSV file with totals has been created: {excel_file_path}")

    
    # for img_tag in soup.find_all('img'):
    #     with open(img_tag['src'], 'rb') as img_file:
    #         img_data = base64.b64encode(img_file.read()).decode('utf-8')
    #     img_tag['src'] = 'data:image/png;base64,' + img_data
    #     # Reduce image size by setting width
    #     img_tag['style'] = 'max-width: 100%; height: auto;'
   
    # for tag in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
    #     tag['style'] = 'font-size: 16px;'  # Increase font size to 16px



    # html_text = str(soup)


    # if project_name:
    #     wrapped_html = add_title_page_to_pdf(html_text, project_name, FFT_LOGO_PATH)
    # else:
    #     wrapped_html = html_text
   
    # print(html_text)
#     pdfkit.from_string(wrapped_html, pdf_file_path, {
#     'page-size': 'Letter',
#     'margin-top': '0.75in',
#     'margin-right': '0.75in',
#     'margin-bottom': '0.75in',
#     'margin-left': '0.75in',
#     # 'footer-center': '[page] of [topage]',
#     # 'footer-left': 'FiftyFive Technologies Pvt Ltd',
#     'enable-local-file-access': None  # Allow access to local files like images
# })

